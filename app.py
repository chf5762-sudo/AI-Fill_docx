import streamlit as st
import json
import os
import base64
import re
import requests
from io import BytesIO
from docx import Document

# API 库导入
from openai import OpenAI, APIConnectionError, AuthenticationError, BadRequestError
import anthropic
import google.generativeai as genai

# ========== 导入提示词库 ==========
try:
    from prompt_library import (
        PROMPT_LIBRARY, 
        DEFAULT_TEMPLATES, 
        GLOBAL_INSTRUCTIONS,
        build_enhanced_prompt
    )
    PROMPT_LIBRARY_AVAILABLE = True
except ImportError:
    PROMPT_LIBRARY_AVAILABLE = False

st.set_page_config(
    page_title="智能文档填充工具",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==========================================
# 🔒 新增：密码保护层 (固定密码: password)
# ==========================================
def check_password():
    """检查密码，如果未通过则停止运行"""
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False

    if not st.session_state.authenticated:
        # 简单的登录界面样式
        st.markdown(
            """
            <style>
            .stTextInput > div > div > input {text-align: center;}
            </style>
            <div style='text-align: center; margin-top: 50px;'>
                <h1>🔒 请输入访问密码</h1>
                <p style='color: gray;'>系统受保护，请输入密码以继续</p>
            </div>
            """, 
            unsafe_allow_html=True
        )
        
        password = st.text_input("密码", type="password", key="login_password")
        
        if password:
            if password == "password": # 这里设置你的固定密码
                st.session_state.authenticated = True
                st.rerun()
            else:
                st.error("❌ 密码错误，请重试")
        
        st.stop() # ⛔ 停止执行后续代码

# 执行密码检查
check_password()

# ==========================================
# ✅ 验证通过，执行主程序
# ==========================================

CONFIG_FILE = "api_config.json"

# ========== API 配置 ==========
API_TYPES = {
    "openai_official": {
        "name": "OpenAI 官方",
        "needs_url": False,
        "default_url": "https://api.openai.com/v1",
        "default_models": ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo"]
    },
    "claude_official": {
        "name": "Claude 官方",
        "needs_url": False,
        "default_url": None,
        "default_models": ["claude-3-5-sonnet-20241022", "claude-3-5-haiku-20241022"]
    },
    "gemini_official": {
        "name": "Gemini 官方",
        "needs_url": False,
        "default_url": None,
        "default_models": ["gemini-2.0-flash-exp", "gemini-1.5-flash", "gemini-1.5-pro"]
    },
    "openai_custom": {
        "name": "OpenAI 自定义",
        "needs_url": True,
        "default_url": "",
        "default_models": []
    },
    "claude_custom": {
        "name": "Claude 自定义",
        "needs_url": True,
        "default_url": "",
        "default_models": []
    },
    "gemini_custom": {
        "name": "Gemini 自定义",
        "needs_url": True,
        "default_url": "",
        "default_models": []
    }
}

# ========== 样式 ==========
st.markdown("""
<style>
    html, body, [class*="css"] { font-size: 16px; }
    .stMarkdown, .stText, p, div, span, label { font-size: 1.1rem !important; }
    .stButton button { font-size: 1.1rem !important; }
    .stTextInput input, .stTextArea textarea, .stSelectbox select { font-size: 1.1rem !important; }
    .stTextInput label, .stTextArea label, .stSelectbox label, .stFileUploader label { font-size: 1.1rem !important; }
    .stAlert, .stInfo, .stWarning, .stSuccess, .stError { font-size: 1.1rem !important; }
    .stCodeBlock, code { font-size: 1rem !important; }
    .streamlit-expanderHeader { font-size: 1.2rem !important; }
    .stTabs [data-baseweb="tab"] { font-size: 1.2rem !important; }
    .main-header { font-size: 2.5rem; font-weight: bold; color: #1f2937; margin-bottom: 0.5rem; }
    .sub-header { font-size: 1.2rem; color: #6b7280; margin-bottom: 2rem; }
    .model-info { background: #f0f9ff; border-left: 4px solid #3b82f6; padding: 0.75rem 1rem; margin: 1rem 0; border-radius: 0.5rem; display: flex; justify-content: space-between; align-items: center; font-size: 1.1rem; }
    .replace-preview { background: #f3f4f6; border-left: 4px solid #3b82f6; padding: 1rem; margin: 0.5rem 0; border-radius: 0.25rem; font-size: 1.1rem; }
    .success-box { background: #d1fae5; border: 2px solid #10b981; padding: 1rem; border-radius: 0.5rem; margin: 1rem 0; font-size: 1.1rem; }
    .error-detail { background: #fee2e2; border: 2px solid #ef4444; padding: 1rem; border-radius: 0.5rem; margin: 1rem 0; font-family: monospace; font-size: 1rem; max-height: 300px; overflow-y: auto; }
</style>
""", unsafe_allow_html=True)

# ========== 配置管理 (已针对云端优化) ==========
def load_config():
    """加载配置文件"""
    default_config = {
        'api_type': 'gemini_custom',
        'api_key': '',
        'base_url': '',
        'model_name': '',
        'model_list': [],
        'prompt_settings': {
            'global_prompt': '',
            'templates': {},
        }
    }
    # 优先从 Streamlit Secrets 读取 (如果配置了)
    if hasattr(st, "secrets") and "api_config" in st.secrets:
        try:
            return {**default_config, **st.secrets["api_config"]}
        except Exception:
            pass

    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                loaded = json.load(f)
                return {**default_config, **loaded}
        except Exception as e:
            st.warning(f"⚠️ 配置加载失败: {str(e)}")
    return default_config

def save_config():
    """保存配置 (增加异常处理，防止云端无权限报错)"""
    config = {
        'api_type': st.session_state.api_type,
        'api_key': st.session_state.api_key,
        'base_url': st.session_state.base_url,
        'model_name': st.session_state.model_name,
        'model_list': st.session_state.model_list,
        'prompt_settings': st.session_state.prompt_settings
    }
    try:
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        return True
    except OSError:
        # 云端环境通常是只读的，无法保存文件，忽略此错误
        return False
    except Exception as e:
        # 其他错误依然提示
        st.error(f"❌ 保存失败: {str(e)}")
        return False

# ========== URL 处理工具 ==========
def get_clean_base_url(url):
    """清洗并标准化 Base URL"""
    if not url:
        return ""
    clean = url.strip().rstrip('/')
    if clean.endswith('/chat/completions'):
        clean = clean.replace('/chat/completions', '')
    if clean.endswith('/models'):
        clean = clean.replace('/models', '')
    if not clean.endswith('/v1'):
        clean += '/v1'
    return clean

# ========== 模型获取功能 ==========
def fetch_models_list(api_type, api_key, base_url):
    """获取模型列表（统一接口）"""
    
    # 官方 OpenAI
    if api_type == "openai_official":
        try:
            client = OpenAI(api_key=api_key, timeout=10)
            models = client.models.list()
            return [m.id for m in models.data if 'gpt' in m.id.lower()]
        except Exception as e:
            return None, f"OpenAI 官方连接失败: {str(e)}"
    
    # 官方 Claude
    elif api_type == "claude_official":
        # Claude 官方不提供模型列表接口，返回预设
        return API_TYPES["claude_official"]["default_models"], None
    
    # 官方 Gemini
    elif api_type == "gemini_official":
        try:
            genai.configure(api_key=api_key)
            models = genai.list_models()
            model_names = [m.name.replace('models/', '') for m in models if 'generateContent' in m.supported_generation_methods]
            return model_names if model_names else API_TYPES["gemini_official"]["default_models"], None
        except Exception as e:
            return None, f"Gemini 官方连接失败: {str(e)}"
    
    # 自定义 API（OpenAI 兼容格式）
    elif api_type in ["openai_custom", "claude_custom", "gemini_custom"]:
        if not base_url:
            return None, "请填写 Base URL"
        
        clean_url = get_clean_base_url(base_url)
        models_url = f"{clean_url}/models"
        
        headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36",
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key if api_key else 'sk-dummy'}"
        }
        
        try:
            response = requests.get(models_url, headers=headers, timeout=10)
            if response.status_code == 200:
                data = response.json()
                if 'data' in data and isinstance(data['data'], list):
                    return [m['id'] for m in data['data']], None
                return None, "返回格式异常"
            else:
                return None, f"HTTP {response.status_code}"
        except Exception as e:
            return None, f"连接失败: {str(e)}"
    
    return None, "未知 API 类型"

# ========== API 测试功能 ==========
def test_api_connection(api_type, api_key, base_url, model_name):
    """测试 API 连接"""
    test_prompt = "请回复：OK"
    
    try:
        response, error = call_ai_api(test_prompt, api_type, api_key, base_url, model_name)
        if error:
            return False, error
        if response and len(response) > 0:
            return True, "连接成功！"
        return False, "返回内容为空"
    except Exception as e:
        return False, str(e)

# ========== 核心 API 调用 ==========
def call_ai_api(prompt, api_type=None, api_key=None, base_url=None, model_name=None, image_data=None, custom_prompt=None):
    """统一的 AI 调用接口"""
    
    # 获取配置
    if api_type is None:
        api_type = st.session_state.get('api_type', 'gemini_custom')
    if api_key is None:
        api_key = st.session_state.get('api_key', '')
    if base_url is None:
        base_url = st.session_state.get('base_url', '')
    if model_name is None:
        model_name = st.session_state.get('model_name', '')
    
    # 增强提示词
    enhanced_prompt = get_enhanced_prompt(prompt, custom_prompt)
    
    try:
        # ========== OpenAI 官方 ==========
        if api_type == "openai_official":
            if not api_key:
                return None, "请配置 OpenAI API Key"
            
            client = OpenAI(api_key=api_key, timeout=60)
            
            if image_data and 'gpt-4' in model_name:
                messages = [{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": enhanced_prompt},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_data}"}}
                    ]
                }]
            else:
                messages = [{"role": "user", "content": enhanced_prompt}]
            
            response = client.chat.completions.create(
                model=model_name,
                messages=messages,
                temperature=0.1
            )
            return response.choices[0].message.content, None
        
        # ========== Claude 官方 ==========
        elif api_type == "claude_official":
            if not api_key:
                return None, "请配置 Claude API Key"
            
            client = anthropic.Anthropic(api_key=api_key)
            
            if image_data:
                content = [
                    {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": image_data}},
                    {"type": "text", "text": enhanced_prompt}
                ]
            else:
                content = enhanced_prompt
            
            message = client.messages.create(
                model=model_name,
                max_tokens=4096,
                messages=[{"role": "user", "content": content}]
            )
            return message.content[0].text, None
        
        # ========== Gemini 官方 ==========
        elif api_type == "gemini_official":
            if not api_key:
                return None, "请配置 Gemini API Key"
            
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(model_name)
            
            if image_data:
                import PIL.Image
                img = PIL.Image.open(BytesIO(base64.b64decode(image_data)))
                response = model.generate_content([enhanced_prompt, img])
            else:
                response = model.generate_content(enhanced_prompt)
            
            return response.text, None
        
        # ========== 自定义 API（OpenAI 兼容格式）==========
        elif api_type in ["openai_custom", "claude_custom", "gemini_custom"]:
            if not base_url:
                return None, "请配置 Base URL"
            
            clean_url = get_clean_base_url(base_url)
            
            client = OpenAI(
                api_key=api_key if api_key else "sk-dummy",
                base_url=clean_url,
                timeout=60.0,
                max_retries=1,
                default_headers={
                    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"
                }
            )
            
            if image_data:
                messages = [{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": enhanced_prompt},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_data}"}}
                    ]
                }]
            else:
                messages = [{"role": "user", "content": enhanced_prompt}]
            
            response = client.chat.completions.create(
                model=model_name,
                messages=messages,
                temperature=0.1
            )
            return response.choices[0].message.content, None
        
        else:
            return None, f"未知的 API 类型: {api_type}"
    
    except AuthenticationError:
        return None, "❌ 认证失败，请检查 API Key"
    except APIConnectionError as e:
        return None, f"❌ 连接失败: {str(e)}"
    except BadRequestError as e:
        return None, f"❌ 请求错误: {str(e)}"
    except Exception as e:
        return None, f"❌ 未知错误: {str(e)}"

# ========== JSON 处理工具 ==========
def clean_json_response(response_text):
    """清理 AI 返回的 JSON 响应"""
    if not response_text:
        return ""
    
    text = response_text.strip()
    
    # 提取 ```json ``` 代码块
    match = re.search(r'```(?:json)?\s*(.*?)\s*```', text, re.DOTALL | re.IGNORECASE)
    if match:
        text = match.group(1)
    
    # 正则提取 JSON 对象
    match = re.search(r'\{(?:[^{}]|(?:\{(?:[^{}]|(?:\{[^{}]*\}))*\}))*\}', text, re.DOTALL)
    if match:
        return match.group(0)
    
    return text

def parse_json_safely(response_text, context=""):
    """安全解析 JSON"""
    try:
        cleaned = clean_json_response(response_text)
        result = json.loads(cleaned)
        
        filtered = {}
        for k, v in result.items():
            if v:
                val = v.get('value', '') if isinstance(v, dict) else str(v)
                if val.strip():
                    filtered[k] = val.strip()
        
        return filtered, None
    except Exception as e:
        error_details = {
            'error': f"JSON 解析失败 ({context}): {str(e)}",
            'original_response': response_text[:500],
            'cleaned_response': clean_json_response(response_text[:500])
        }
        return None, error_details

# ========== 提示词增强 ==========
def get_enhanced_prompt(base_prompt, custom_prompt=None):
    """增强提示词"""
    if custom_prompt and custom_prompt.strip():
        return f"{custom_prompt}\n\n{base_prompt}"
    
    if not PROMPT_LIBRARY_AVAILABLE:
        return base_prompt
    
    prompt_settings = st.session_state.get('prompt_settings', {})
    enhanced = base_prompt
    
    templates = prompt_settings.get('templates', {})
    for name, template_config in templates.items():
        if template_config.get('enabled', False):
            if t_prompt := template_config.get('prompt', '').strip():
                enhanced = f"【文档类型: {name}】\n{t_prompt}\n\n{enhanced}"
            break
    
    if g_prompt := prompt_settings.get('global_prompt', '').strip():
        enhanced = f"【全局指令】\n{g_prompt}\n\n{enhanced}"
    
    return enhanced

# ========== 业务逻辑函数 ==========
def extract_customer_info_from_text(text, image_data=None, custom_prompt=None):
    """从用户输入中提取信息"""
    prompt = f"""
你是信息提取专家，从混乱的文本中全面识别信息。

【用户输入】
{text}

【识别规则】
请提取以下所有能找到的信息：
1. 公司名称、联系人姓名、职位/部门
2. 联系电话、手机号码、固定电话、传真号码
3. 电子邮箱、详细地址、邮政编码、公司网址
4. 统一社会信用代码、法人代表、开户行信息、银行账号
5. 任何其他看起来重要的信息

【输出格式】
只返回纯 JSON 对象，不要任何其他文字：
{{
  "公司名称": "提取的值",
  "联系人": "提取的值",
  ...
}}
"""
    
    response_text, error = call_ai_api(prompt, image_data=image_data, custom_prompt=custom_prompt)
    if error:
        return None, error
    return parse_json_safely(response_text, "提取客户信息")

def extract_document_content(doc):
    """提取 Word 文档内容"""
    content_parts = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            content_parts.append(f"[P{i}] {para.text.strip()}")
    
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            row_data = [c.text.strip() for c in row.cells]
            if any(row_data):
                content_parts.append(f"[T{ti}-R{ri}] {' | '.join(row_data)}")
    
    return "\n".join(content_parts)

def analyze_reference_document(doc, custom_prompt=None):
    """分析参考文档"""
    st.info("📄 正在提取文档内容...")
    doc_content = extract_document_content(doc)
    
    if len(doc_content) > 15000:
        doc_content = doc_content[:15000] + "\n...(内容过长，已截断)"
    
    st.info("🤖 AI 正在全面分析可变数据...")
    
    prompt = f"""
这是一份已填写好的文档内容。请提取所有"可变的、会随具体情况变化"的数据。

【文档内容】
{doc_content}

【任务】
提取所有会随不同情况变化的信息，如：公司名称、联系人、电话、地址、型号、数量、金额、日期等。

【输出格式】
只返回纯 JSON 对象：
{{
  "公司名称": "具体值",
  "联系人": "具体值",
  ...
}}
"""
    
    response_text, error = call_ai_api(prompt, custom_prompt=custom_prompt)
    if error:
        return None, error
    return parse_json_safely(response_text, "分析参考文档")

def create_replacement_mapping(old_info, new_info, custom_prompt=None):
    """创建替换映射"""
    prompt = f"""
【参考数据（旧）】
{json.dumps(old_info, ensure_ascii=False, indent=2)}

【新数据】
{json.dumps(new_info, ensure_ascii=False, indent=2)}

【任务】
建立新旧数据的对应关系。如果新数据缺失则对应 null。

【输出格式】
只返回纯 JSON 对象（值对值的映射）：
{{
  "旧值1": "新值1",
  "旧值2": "新值2",
  "旧值3": null
}}
"""
    
    response_text, error = call_ai_api(prompt, custom_prompt=custom_prompt)
    if error:
        return None, error
    return parse_json_safely(response_text, "创建替换映射")

# =======================================================
# 📌 修复：Word 文档替换逻辑
# =======================================================

def replace_text_in_paragraph(paragraph, old_text, new_text):
    """
    【简化强化版】更可靠的段落文本替换
    策略：完全重建段落 runs，保留第一个 run 的格式
    
    返回: (replace_success: bool, count: int)
    """
    full_text = paragraph.text
    
    if old_text not in full_text:
        return False, 0
    
    # 执行一次替换
    new_full_text = full_text.replace(old_text, new_text, 1)
    
    if new_full_text == full_text:
        return False, 0
    
    # 保存第一个 run 的格式（如果存在）
    first_run_format = None
    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run_format = {
            'bold': first_run.bold,
            'italic': first_run.italic,
            'underline': first_run.underline,
            'font_name': first_run.font.name,
            'font_size': first_run.font.size,
        }
    
    # 清空所有 runs
    for run in paragraph.runs:
        run.text = ""
    
    # 创建新 run 或使用第一个 run
    if paragraph.runs:
        new_run = paragraph.runs[0]
    else:
        new_run = paragraph.add_run()
    
    new_run.text = new_full_text
    
    # 恢复格式
    if first_run_format:
        try:
            new_run.bold = first_run_format['bold']
            new_run.italic = first_run_format['italic']
            new_run.underline = first_run_format['underline']
            if first_run_format['font_name']:
                new_run.font.name = first_run_format['font_name']
            if first_run_format['font_size']:
                new_run.font.size = first_run_format['font_size']
        except:
            pass  # 格式恢复失败不影响替换
    
    return True, 1


def apply_replacements_to_document(doc, replacement_mapping, debug_mode=False):
    """应用替换到文档（增加调试信息）"""
    replace_count = 0
    replace_log = []
    debug_info = []
    
    # 按长度排序，避免短字符串误替换
    sorted_map = sorted(replacement_mapping.items(), key=lambda x: len(str(x[0])), reverse=True)
    
    for old_val, new_val in sorted_map:
        if not old_val or not new_val:
            continue
        
        current_count = 0
        old_s = str(old_val)
        new_s = str(new_val)
        
        if debug_mode:
            debug_info.append(f"\n🔍 开始处理: '{old_s}' → '{new_s}'")
        
        # 1. 替换段落
        para_count = 0
        for idx, p in enumerate(doc.paragraphs):
            if old_s in p.text:
                if debug_mode:
                    debug_info.append(f"  📝 段落 {idx} 包含目标文本")
                    debug_info.append(f"     原文: {p.text[:100]}")
                
                total_paragraph_replacements = 0
                while True:
                    success, count = replace_text_in_paragraph(p, old_s, new_s)
                    total_paragraph_replacements += count
                    if not success:
                        break
                
                if debug_mode and total_paragraph_replacements > 0:
                    debug_info.append(f"     ✅ 替换了 {total_paragraph_replacements} 处")
                    debug_info.append(f"     结果: {p.text[:100]}")
                
                para_count += total_paragraph_replacements
        
        current_count += para_count
        
        # 2. 替换表格
        table_count = 0
        for ti, t in enumerate(doc.tables):
            for ri, r in enumerate(t.rows):
                for ci, c in enumerate(r.cells):
                    for pi, p in enumerate(c.paragraphs):
                        if old_s in p.text:
                            if debug_mode:
                                debug_info.append(f"  📊 表格 {ti} 行 {ri} 列 {ci} 包含目标文本")
                            
                            total_cell_replacements = 0
                            while True:
                                success, count = replace_text_in_paragraph(p, old_s, new_s)
                                total_cell_replacements += count
                                if not success:
                                    break
                            
                            if debug_mode and total_cell_replacements > 0:
                                debug_info.append(f"     ✅ 替换了 {total_cell_replacements} 处")
                            
                            table_count += total_cell_replacements
        
        current_count += table_count
        
        # 3. 替换页眉页脚
        header_count = 0
        for si, section in enumerate(doc.sections):
            for header in [section.header, section.footer]:
                if header:
                    for pi, p in enumerate(header.paragraphs):
                        if old_s in p.text:
                            if debug_mode:
                                debug_info.append(f"  📋 节 {si} 页眉/页脚包含目标文本")
                            
                            total_header_replacements = 0
                            while True:
                                success, count = replace_text_in_paragraph(p, old_s, new_s)
                                total_header_replacements += count
                                if not success:
                                    break
                            
                            if debug_mode and total_header_replacements > 0:
                                debug_info.append(f"     ✅ 替换了 {total_header_replacements} 处")
                            
                            header_count += total_header_replacements
        
        current_count += header_count
        
        if current_count > 0:
            replace_count += current_count
            replace_log.append(f"✓ 替换 '{old_s}' → '{new_s}' ({current_count}处: 段落{para_count}+表格{table_count}+页眉页脚{header_count})")
        else:
            replace_log.append(f"⚠️ 未找到 '{old_s}'")
            if debug_mode:
                debug_info.append(f"  ❌ 全文未找到此文本")
    
    if debug_mode:
        return replace_count, replace_log, debug_info
    
    return replace_count, replace_log

# =======================================================
# 📌 修复结束
# =======================================================


# ========== 初始化 Session State ==========
if 'step' not in st.session_state:
    st.session_state.step = 1
if 'show_prompt_editor' not in st.session_state:
    st.session_state.show_prompt_editor = False

for k in ['template_file', 'template_filename', 'old_customer_info', 'new_customer_info', 
          'replacement_mapping', 'uploaded_image_data', 'custom_replacements', 'current_prompt']:
    if k not in st.session_state:
        st.session_state[k] = None if 'file' in k or 'image' in k or 'prompt' in k else {}

if st.session_state.custom_replacements is None:
    st.session_state.custom_replacements = []

# 加载配置
cfg = load_config()
if 'api_type' not in st.session_state:
    st.session_state.api_type = cfg.get('api_type', 'gemini_custom')
if 'api_key' not in st.session_state:
    st.session_state.api_key = cfg.get('api_key', '')
if 'base_url' not in st.session_state:
    st.session_state.base_url = cfg.get('base_url', '')
if 'model_name' not in st.session_state:
    st.session_state.model_name = cfg.get('model_name', '')
if 'model_list' not in st.session_state:
    st.session_state.model_list = cfg.get('model_list', [])
if 'prompt_settings' not in st.session_state:
    st.session_state.prompt_settings = cfg.get('prompt_settings', {})

# ==================== 侧边栏 ====================
with st.sidebar:
    st.markdown("## ⚙️ API 配置")
    
    # 1. 选择 API 类型
    api_type_options = list(API_TYPES.keys())
    api_type_labels = [API_TYPES[k]["name"] for k in api_type_options]
    
    current_index = api_type_options.index(st.session_state.api_type) if st.session_state.api_type in api_type_options else 0
    
    selected_label = st.selectbox(
        "API 类型",
        options=api_type_labels,
        index=current_index,
        key="api_type_selector"
    )
    
    selected_type = api_type_options[api_type_labels.index(selected_label)]
    
    if selected_type != st.session_state.api_type:
        st.session_state.api_type = selected_type
        st.session_state.model_list = []
        st.session_state.model_name = ''
        save_config()
        st.rerun()
    
    # 2. API Key
    api_key_input = st.text_input(
        "API Key" + (" *必填" if "official" in st.session_state.api_type else " (可选)"),
        value=st.session_state.api_key,
        type="password",
        key="api_key_input"
    )
    
    if api_key_input != st.session_state.api_key:
        st.session_state.api_key = api_key_input
        save_config()
    
    # 3. Base URL（仅自定义需要）
    if API_TYPES[st.session_state.api_type]["needs_url"]:
        base_url_input = st.text_input(
            "Base URL *必填",
            value=st.session_state.base_url,
            placeholder="https://xxx.workers.dev",
            key="base_url_input"
        )
        
        if base_url_input != st.session_state.base_url:
            st.session_state.base_url = base_url_input
            save_config()
    
    st.markdown("---")
    
    # 4. 模型管理
    st.markdown("### 📋 模型管理")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("🔄 获取模型", use_container_width=True):
            with st.spinner("正在获取模型列表..."):
                models, error = fetch_models_list(
                    st.session_state.api_type,
                    st.session_state.api_key,
                    st.session_state.base_url
                )
                
                if error:
                    st.error(f"❌ {error}")
                elif models:
                    st.session_state.model_list = models
                    if not st.session_state.model_name or st.session_state.model_name not in models:
                        st.session_state.model_name = models[0]
                    save_config()
                    st.success(f"✅ 获取成功！共 {len(models)} 个模型")
                    st.rerun()
                else:
                    st.warning("未获取到模型列表")
    
    with col2:
        if st.button("🧪 测试连接", use_container_width=True):
            if not st.session_state.model_name:
                st.warning("请先选择模型")
            else:
                with st.spinner("正在测试..."):
                    success, message = test_api_connection(
                        st.session_state.api_type,
                        st.session_state.api_key,
                        st.session_state.base_url,
                        st.session_state.model_name
                    )
                    
                    if success:
                        st.success(message)
                    else:
                        st.error(message)
    
    # 5. 模型选择
    if st.session_state.model_list:
        model_options = list(st.session_state.model_list)
        if st.session_state.model_name and st.session_state.model_name not in model_options:
            model_options.insert(0, st.session_state.model_name)
        
        current_model_index = model_options.index(st.session_state.model_name) if st.session_state.model_name in model_options else 0
        
        selected_model = st.selectbox(
            "选择模型",
            options=model_options,
            index=current_model_index,
            key="model_selector"
        )
        
        if selected_model != st.session_state.model_name:
            st.session_state.model_name = selected_model
            save_config()
            st.rerun()
    else:
        # 手动输入模型名称
        model_input = st.text_input(
            "模型名称（手动输入）",
            value=st.session_state.model_name,
            placeholder="例如: gemini-1.5-flash",
            key="model_name_input"
        )
        
        if model_input != st.session_state.model_name:
            st.session_state.model_name = model_input
            save_config()
    
    st.markdown("---")
    
    # 6. 格式说明
    st.markdown("## 📄 格式说明")
    st.info("""
**仅支持 .docx 格式**

**转换工具推荐：**
• [Smallpdf](https://smallpdf.com/cn/pdf-to-word)  
• [ILovePDF](https://www.ilovepdf.com/zh-cn/pdf_to_word)  
• [Convertio](https://convertio.co/zh/pdf-docx/)
    """)

# ==================== 主界面 ====================
st.markdown('<div class="main-header">📄 智能文档填充工具</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">仿照模式 - AI学习已填好的文档</div>', unsafe_allow_html=True)

# 顶部模型信息
if st.session_state.api_key or st.session_state.base_url:
    api_name = API_TYPES.get(st.session_state.api_type, {}).get("name", "未知")
    st.markdown(f"""
    <div class="model-info">
        <span>✅ {api_name} | 模型: <code>{st.session_state.model_name or '未选择'}</code></span>
    </div>
    """, unsafe_allow_html=True)
else:
    st.warning("⚠️ 请在左侧侧边栏配置 API")

# 进度指示
progress_cols = st.columns(5)
steps = ["上传文档", "AI分析", "输入数据", "确认替换", "下载"]
for i, col in enumerate(progress_cols, 1):
    with col:
        if st.session_state.step == i:
            st.markdown(f"### ✅ {steps[i-1]}")
        elif st.session_state.step > i:
            st.markdown(f"### ✓ {steps[i-1]}")
        else:
            st.markdown(f"### ⭕ {steps[i-1]}")

st.markdown("---")

# ==================== 步骤1: 上传参考文档 ====================
if st.session_state.step >= 1:
    st.markdown("## 步骤1️⃣: 上传参考文档")
    st.info("💡 上传一份已经填写好的文档，AI会学习它的填写方式")
    
    uploaded_file = st.file_uploader(
        "选择已填好的文档（.docx）",
        type=['docx'],
        help="Word 文档格式"
    )
    
    if uploaded_file:
        st.session_state.template_file = uploaded_file
        st.session_state.template_filename = uploaded_file.name
        st.success(f"✅ 已上传: {uploaded_file.name}")
        
        if st.button("下一步：AI分析 ➡️", type="primary"):
            st.session_state.step = 2
            st.rerun()

# ==================== 步骤2: AI分析文档 ====================
if st.session_state.step >= 2:
    st.markdown("## 步骤2️⃣: AI分析参考文档")
    
    # 提示词编辑器
    if st.session_state.step == 2:
        with st.expander("💡 查看/编辑提示词（可选）", expanded=st.session_state.show_prompt_editor):
            st.markdown("### 临时自定义提示词")
            st.caption("仅在本次分析中生效")
            
            custom_prompt = st.text_area(
                "自定义提示词",
                value=st.session_state.current_prompt or "",
                height=200,
                placeholder="例如：\n• 重点识别技术参数和型号\n• 忽略通用条款\n• 优先提取数量和金额信息",
                key="custom_prompt_analyze"
            )
            
            col1, col2 = st.columns(2)
            with col1:
                if st.button("✅ 应用此提示词", use_container_width=True, type="primary"):
                    st.session_state.current_prompt = custom_prompt.strip() if custom_prompt.strip() else None
                    st.success("✅ 已应用临时提示词")
                    st.session_state.show_prompt_editor = False
                    st.rerun()
            
            with col2:
                if st.button("🔄 恢复默认", use_container_width=True):
                    st.session_state.current_prompt = None
                    st.success("✅ 已恢复默认配置")
                    st.rerun()
    
    if st.session_state.step == 2 and not st.session_state.old_customer_info:
        if not st.session_state.api_key and "official" in st.session_state.api_type:
            st.error("❌ 请先在侧边栏配置 API Key")
        elif not st.session_state.model_name:
            st.error("❌ 请先在侧边栏选择模型")
        else:
            # 关键修复：指针复位
            st.session_state.template_file.seek(0)
            doc = Document(st.session_state.template_file)
            
            old_info, error = analyze_reference_document(doc, st.session_state.current_prompt)
            
            if error:
                st.error("❌ 分析失败")
                
                with st.expander("🔍 查看详细错误信息", expanded=True):
                    st.markdown('<div class="error-detail">', unsafe_allow_html=True)
                    if isinstance(error, dict):
                        st.markdown(f"**错误类型:** {error.get('error', '未知错误')}")
                        if 'original_response' in error:
                            st.markdown("**原始返回内容:**")
                            st.code(error['original_response'], language='text')
                        if 'cleaned_response' in error:
                            st.markdown("**清理后内容:**")
                            st.code(error['cleaned_response'], language='text')
                    else:
                        st.code(str(error))
                    st.markdown('</div>', unsafe_allow_html=True)
                
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("⬅️ 返回重试", use_container_width=True):
                        st.session_state.step = 1
                        st.session_state.old_customer_info = {}
                        st.session_state.current_prompt = None
                        st.rerun()
                with col2:
                    if st.button("💡 调整提示词", use_container_width=True):
                        st.session_state.show_prompt_editor = True
                        st.rerun()
            else:
                st.session_state.old_customer_info = old_info
                st.success("✅ 分析完成！")
                st.rerun()
    
    if st.session_state.old_customer_info:
        st.markdown('<div class="success-box">', unsafe_allow_html=True)
        st.markdown("### 识别到的可变数据：")
        cols = st.columns(2)
        for idx, (field, value) in enumerate(st.session_state.old_customer_info.items()):
            with cols[idx % 2]:
                st.markdown(f"**{field}:** `{value}`")
        st.markdown('</div>', unsafe_allow_html=True)
        
        col1, col2 = st.columns([1, 3])
        with col1:
            if st.button("⬅️ 重新上传"):
                st.session_state.step = 1
                st.session_state.old_customer_info = {}
                st.session_state.current_prompt = None
                st.rerun()
        with col2:
            if st.button("下一步：输入新数据 ➡️", type="primary", use_container_width=True):
                st.session_state.step = 3
                st.rerun()

# ==================== 步骤3: 输入新数据 ====================
if st.session_state.step >= 3:
    st.markdown("## 步骤3️⃣: 输入新数据")
    st.info("💡 随意输入，AI会自动识别格式")
    
    # 提示词编辑器
    with st.expander("💡 查看/编辑提示词（可选）"):
        st.markdown("### 临时自定义提示词")
        st.caption("仅在本次提取中生效")
        
        custom_prompt_extract = st.text_area(
            "自定义提示词",
            value="",
            height=150,
            placeholder="例如：\n• 重点识别联系方式\n• 提取所有电话和邮箱",
            key="custom_prompt_extract"
        )
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📝 文本输入")
        input_text = st.text_area(
            "新数据资料",
            height=300,
            placeholder="""例如：
上海普宙科技
张经理
mobile: 15912345678
浦东新区张江

或者：
Company: Shanghai Puzhou
Contact: Manager Zhang  
Tel 159-1234-5678"""
        )
    
    with col2:
        st.markdown("### 📷 图片识别")
        uploaded_image = st.file_uploader(
            "上传图片（名片/截图等）",
            type=['jpg', 'jpeg', 'png']
        )
        if uploaded_image:
            st.image(uploaded_image, use_container_width=True)
            image_bytes = uploaded_image.read()
            st.session_state.uploaded_image_data = base64.b64encode(image_bytes).decode()
            uploaded_image.seek(0)
    
    col_btn1, col_btn2 = st.columns([1, 1])
    
    with col_btn1:
        if st.button("🤖 AI提取", type="primary", use_container_width=True):
            if not input_text and not st.session_state.uploaded_image_data:
                st.warning("请输入文本或上传图片")
            else:
                with st.spinner("AI分析中..."):
                    custom_p = custom_prompt_extract.strip() if custom_prompt_extract.strip() else None
                    new_info, error = extract_customer_info_from_text(
                        input_text, 
                        st.session_state.uploaded_image_data,
                        custom_p
                    )
                    
                    if error:
                        st.error("❌ 提取失败")
                        
                        with st.expander("🔍 查看详细错误信息", expanded=True):
                            st.markdown('<div class="error-detail">', unsafe_allow_html=True)
                            if isinstance(error, dict):
                                st.markdown(f"**错误类型:** {error.get('error', '未知错误')}")
                                if 'original_response' in error:
                                    st.markdown("**原始返回内容:**")
                                    st.code(error['original_response'], language='text')
                                if 'cleaned_response' in error:
                                    st.markdown("**清理后内容:**")
                                    st.code(error['cleaned_response'], language='text')
                            else:
                                st.code(str(error))
                            st.markdown('</div>', unsafe_allow_html=True)
                    else:
                        new_info = {k: v for k, v in new_info.items() if v and v.strip()}
                        st.session_state.new_customer_info = new_info
                        st.success("✅ 提取完成")
                        st.session_state.step = 4
                        st.rerun()
    
    with col_btn2:
        if st.button("⬅️ 返回", use_container_width=True):
            st.session_state.step = 2
            st.rerun()

# ==================== 步骤4: 确认替换映射 ====================
if st.session_state.step >= 4:
    st.markdown("## 步骤4️⃣: 确认替换内容")
    
    if st.session_state.step == 4 and not st.session_state.replacement_mapping:
        with st.spinner("AI正在匹配新旧数据..."):
            mapping, error = create_replacement_mapping(
                st.session_state.old_customer_info,
                st.session_state.new_customer_info
            )
            
            if error:
                st.error("❌ 匹配失败")
                
                with st.expander("🔍 查看详细错误信息", expanded=True):
                    st.markdown('<div class="error-detail">', unsafe_allow_html=True)
                    if isinstance(error, dict):
                        st.markdown(f"**错误类型:** {error.get('error', '未知错误')}")
                        if 'original_response' in error:
                            st.markdown("**原始返回内容:**")
                            st.code(error['original_response'], language='text')
                        if 'cleaned_response' in error:
                            st.markdown("**清理后内容:**")
                            st.code(error['cleaned_response'], language='text')
                    else:
                        st.code(str(error))
                    st.markdown('</div>', unsafe_allow_html=True)
                
                if st.button("⬅️ 返回"):
                    st.session_state.step = 3
                    st.session_state.replacement_mapping = {}
                    st.rerun()
            else:
                st.session_state.replacement_mapping = mapping
                st.success("✅ 匹配完成")
    
    if st.session_state.replacement_mapping:
        st.markdown("### 📋 AI识别的替换项")
        st.info("💡 可以编辑新值，或取消某项替换")
        
        edited_mapping = {}
        
        # 排除 null 值，只展示有替换意向的项
        filtered_mapping = {k: v for k, v in st.session_state.replacement_mapping.items() if v is not None}

        for old_val, new_val in filtered_mapping.items():
            st.markdown('<div class="replace-preview">', unsafe_allow_html=True)
            
            col1, col2, col3 = st.columns([2, 1, 2])
            
            with col1:
                st.markdown("**旧值:**")
                st.code(old_val)
            
            with col2:
                st.markdown("**→**")
            
            with col3:
                st.markdown("**新值:**")
                edited_val = st.text_input(
                    f"edit_{old_val}",
                    value=new_val,
                    label_visibility="collapsed",
                    key=f"edit_{hash(old_val)}"
                )
                edited_mapping[old_val] = edited_val
            
            st.markdown('</div>', unsafe_allow_html=True)
        
        st.session_state.replacement_mapping = edited_mapping
        
        st.markdown("---")
        
        st.markdown("### ➕ 手动添加替换项")
        with st.expander("添加自定义替换", expanded=False):
            st.info("💡 如果AI未识别到某些需要替换的内容，可以手动添加")
            
            if st.session_state.custom_replacements:
                st.markdown("**已添加的自定义项：**")
                items_to_remove = []
                for idx, item in enumerate(st.session_state.custom_replacements):
                    col_display, col_delete = st.columns([9, 1])
                    with col_display:
                        st.markdown(f"• `{item['old']}` → `{item['new']}`")
                    with col_delete:
                        if st.button("🗑️", key=f"del_custom_{idx}", help="删除"):
                            items_to_remove.append(idx)
                
                if items_to_remove:
                    for idx in sorted(items_to_remove, reverse=True):
                        item = st.session_state.custom_replacements[idx]
                        if item['old'] in st.session_state.replacement_mapping:
                            del st.session_state.replacement_mapping[item['old']]
                        st.session_state.custom_replacements.pop(idx)
                    st.rerun()
                
                st.markdown("---")
            
            col_old, col_new, col_add = st.columns([5, 5, 2])
            
            with col_old:
                custom_old = st.text_input(
                    "旧值",
                    placeholder="例如：旧公司名",
                    key="custom_old_input"
                )
            
            with col_new:
                custom_new = st.text_input(
                    "新值",
                    placeholder="例如：新公司名",
                    key="custom_new_input"
                )
            
            with col_add:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("➕ 添加", type="primary", use_container_width=True):
                    if custom_old and custom_new:
                        already_exists = False
                        for item in st.session_state.custom_replacements:
                            if item['old'] == custom_old.strip():
                                already_exists = True
                                break
                        
                        if already_exists:
                            st.warning("⚠️ 该旧值已存在")
                        else:
                            st.session_state.replacement_mapping[custom_old.strip()] = custom_new.strip()
                            st.session_state.custom_replacements.append({
                                'old': custom_old.strip(),
                                'new': custom_new.strip()
                            })
                            st.success("✅ 已添加")
                            st.rerun()
                    else:
                        st.warning("请填写旧值和新值")
        
        st.markdown("---")
        
        col1, col2 = st.columns([1, 3])
        with col1:
            if st.button("⬅️ 重新输入"):
                st.session_state.step = 3
                st.session_state.replacement_mapping = {}
                st.session_state.custom_replacements = []
                st.rerun()
        
        with col2:
            if st.button("✅ 确认并生成文档", type="primary", use_container_width=True):
                st.session_state.step = 5
                st.rerun()

# ==================== 步骤5: 生成并下载文档 ====================
if st.session_state.step >= 5:
    st.markdown("## 步骤5️⃣: 生成新文档")
    
    # 添加调试模式开关
    debug_mode = st.checkbox("🔍 启用调试模式（查看详细替换过程）", value=False)
    
    with st.spinner("正在生成新文档..."):
        # 关键修复：指针复位
        st.session_state.template_file.seek(0)
        doc = Document(st.session_state.template_file)
        
        if debug_mode:
            replace_count, replace_log, debug_info = apply_replacements_to_document(
                doc, 
                st.session_state.replacement_mapping,
                debug_mode=True
            )
            
            # 显示调试信息
            with st.expander("🐛 调试信息", expanded=True):
                st.code("\n".join(debug_info), language="text")
        else:
            replace_count, replace_log = apply_replacements_to_document(
                doc, 
                st.session_state.replacement_mapping
            )
        
        # 关键修复：使用 BytesIO 内存操作
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        if replace_count > 0:
            st.success(f"✅ 文档生成完成！共替换 {replace_count} 处")
        else:
            st.warning(f"⚠️ 文档已生成，但未执行任何替换。请检查原文档中是否包含待替换内容。")
        
        with st.expander("📋 查看替换详情", expanded=True):
            for log in replace_log:
                st.markdown(log)
        
        original_name = st.session_state.template_filename
        # 修复文件名后缀
        new_filename = f"已填充_{original_name}"
        if not new_filename.lower().endswith(".docx"):
            new_filename = original_name.replace('.docx', '') + '_已填充.docx'
        
        st.download_button(
            label="⬇️ 下载新文档",
            data=output,
            file_name=new_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )
        
        st.markdown("---")
        
        if st.button("🔄 重新开始", use_container_width=True):
            st.session_state.step = 1
            st.session_state.template_file = None
            st.session_state.template_filename = ''
            st.session_state.old_customer_info = {}
            st.session_state.new_customer_info = {}
            st.session_state.replacement_mapping = {}
            st.session_state.uploaded_image_data = None
            st.session_state.custom_replacements = []
            st.session_state.current_prompt = None
            st.rerun()

# ==================== 页脚 ====================
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #6b7280; padding: 2rem 0;'>
    <p>💡 提示：所有配置保存在本地 <code>api_config.json</code> 文件中</p>
    <p>🔒 支持官方和自定义API，保护您的数据隐私</p>
</div>
""", unsafe_allow_html=True)
